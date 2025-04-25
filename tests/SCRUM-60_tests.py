import os
import shutil
import tempfile
import pytest

from requirement_extractor import RequirementExtractor, ExtractionError

def create_sample_pdf(path):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)
    pdf.cell(200, 10, txt="Requirement: The system shall support user login.", ln=True)
    pdf.output(path)

def create_sample_docx(path):
    from docx import Document
    doc = Document()
    doc.add_heading('Requirements', 0)
    doc.add_paragraph('Requirement: The system shall support password reset.')
    doc.save(path)

def create_sample_xlsx(path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws['A1'] = 'Requirement'
    ws['A2'] = 'The system shall send notifications.'
    wb.save(path)

def create_sample_image(path):
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new('RGB', (400, 100), color = (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10, 40), "Requirement: The system shall log all actions.", fill=(0,0,0))
    img.save(path)

@pytest.fixture(scope="module")
def temp_dir():
    d = tempfile.mkdtemp()
    yield d
    shutil.rmtree(d)

@pytest.fixture(scope="module")
def extractor():
    return RequirementExtractor()

def test_extract_requirements_from_pdf(temp_dir, extractor):
    pdf_path = os.path.join(temp_dir, 'sample.pdf')
    create_sample_pdf(pdf_path)
    requirements = extractor.extract(pdf_path)
    assert isinstance(requirements, list)
    assert any('login' in req.lower() for req in requirements)

def test_extract_requirements_from_docx(temp_dir, extractor):
    docx_path = os.path.join(temp_dir, 'sample.docx')
    create_sample_docx(docx_path)
    requirements = extractor.extract(docx_path)
    assert isinstance(requirements, list)
    assert any('password reset' in req.lower() for req in requirements)

def test_extract_requirements_from_xlsx(temp_dir, extractor):
    xlsx_path = os.path.join(temp_dir, 'sample.xlsx')
    create_sample_xlsx(xlsx_path)
    requirements = extractor.extract(xlsx_path)
    assert isinstance(requirements, list)
    assert any('notifications' in req.lower() for req in requirements)

def test_extract_requirements_from_image(temp_dir, extractor):
    image_path = os.path.join(temp_dir, 'sample.png')
    create_sample_image(image_path)
    requirements = extractor.extract(image_path)
    assert isinstance(requirements, list)
    assert any('log all actions' in req.lower() for req in requirements)

def test_extract_requirements_from_unsupported_format(temp_dir, extractor):
    txt_path = os.path.join(temp_dir, 'sample.txt')
    with open(txt_path, 'w') as f:
        f.write('Requirement: The system shall be robust.')
    with pytest.raises(ExtractionError):
        extractor.extract(txt_path)

def test_extract_requirements_from_corrupted_pdf(temp_dir, extractor):
    pdf_path = os.path.join(temp_dir, 'corrupted.pdf')
    with open(pdf_path, 'wb') as f:
        f.write(b'%PDF-1.4 corrupted content')
    with pytest.raises(ExtractionError):
        extractor.extract(pdf_path)

def test_extract_requirements_from_empty_file(temp_dir, extractor):
    empty_path = os.path.join(temp_dir, 'empty.docx')
    with open(empty_path, 'wb') as f:
        pass
    with pytest.raises(ExtractionError):
        extractor.extract(empty_path)

def test_extract_requirements_with_no_requirements(temp_dir, extractor):
    docx_path = os.path.join(temp_dir, 'noreq.docx')
    from docx import Document
    doc = Document()
    doc.add_paragraph('This document contains no requirements.')
    doc.save(docx_path)
    requirements = extractor.extract(docx_path)
    assert isinstance(requirements, list)
    assert len(requirements) == 0
